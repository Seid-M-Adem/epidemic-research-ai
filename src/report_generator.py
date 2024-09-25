from fpdf import FPDF
from docx import Document

def generate_pdf_paper(title, introduction, methods, results, discussion, conclusion, references):
    pdf = FPDF()
    pdf.add_page()
    
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(200, 10, title, ln=True, align='C')
    
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(0, 10, f"Introduction: {introduction}\n\n")
    pdf.multi_cell(0, 10, f"Methods: {methods}\n\n")
    pdf.multi_cell(0, 10, f"Results: {results}\n\n")
    pdf.multi_cell(0, 10, f"Discussion: {discussion}\n\n")
    pdf.multi_cell(0, 10, f"Conclusion: {conclusion}\n\n")
    pdf.multi_cell(0, 10, f"References:\n{references}")
    
    pdf.output('/workspaces/epidemic-research-ai/reports/full_paper.pdf')

def generate_word_paper(title, introduction, methods, results, discussion, conclusion, references):
    doc = Document()
    
    doc.add_heading(title, 0)
    
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(introduction)
    
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(methods)
    
    doc.add_heading('Results', level=1)
    doc.add_paragraph(results)
    
    doc.add_heading('Discussion', level=1)
    doc.add_paragraph(discussion)
    
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(conclusion)
    
    doc.add_heading('References', level=1)
    doc.add_paragraph(references)
    
    doc.save('/workspaces/epidemic-research-ai/reports/full_paper.docx')
